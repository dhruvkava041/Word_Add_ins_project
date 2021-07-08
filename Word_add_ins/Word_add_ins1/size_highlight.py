import os
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.enum.text import WD_COLOR_INDEX


def txt_highlight(filename):
    doc=Document(filename)#Document object
    '''
    para=doc.paragraphs
    runli=para[0].runs#list containing run objects
    print(runli[0].font.size.pt)
    print(runli[1].font.size.pt)
    print(runli[0].text)
    print(type(runli[1].text))
    print(runli[2].text)
    '''
  
    for para in doc.paragraphs:
        for runli in para.runs:
            if (int(runli.font.size.pt))!=12:
                runli.font.highlight_color =WD_COLOR_INDEX.YELLOW

    doc.save("app_output\\highlight_output.docx")
    os.system("start app_output\\highlight_output.docx")
    


txt_highlight("app_input\\text_size.docx")
