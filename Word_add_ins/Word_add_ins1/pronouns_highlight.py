import os
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.enum.text import WD_COLOR_INDEX

def pronouns_highlight(filename):
    doc=Document(filename)
    pronoun_list=["I","me"]
    #for pronoun in pronoun_list:
    for para in doc.paragraphs:
        #for runli in para.runs:
            #if(pronoun in para.text):
                textstring=para.text
                para.clear()
                textlist=textstring.split()
                print(textlist)
                for word in textlist:
                    if word in pronoun_list:
                        new_run=para.add_run(word+" ")
                        new_run.font.highlight_color=WD_COLOR_INDEX.YELLOW
                    else:
                        new_run=para.add_run(word+" ")
                        new_run.font.highlight_color=0

                    
    
    doc.save("app_output\\pronouns_highlight_output.docx")
    os.system("start app_output\\pronouns_highlight_output.docx")

pronouns_highlight("app_input\\test_pronoun.docx")



