import os
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.enum.text import WD_COLOR_INDEX
from termcolor import colored
import copy
import re

def coding(filename):
    doc = Document(filename)
    negativList = ["I","me"]
    for paragraph in doc.paragraphs:
        for target in negativList:
            if target in paragraph.text:  # it is worth checking in detail ...

                currRuns = copy.copy(paragraph.runs)   # deep copy as we delete/clear the object
                #paragraph.runs.clear()
                paragraph.clear()
                for run in currRuns:
                    if target in run.text:
                        words = re.split('(\W)', run.text)  # split into words in order to be able to color only one
                        for word in words:
                            if word == target:
                                newRun = paragraph.add_run(word)
                                newRun.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            else:
                                newRun = paragraph.add_run(word)
                                newRun.font.highlight_color = None
                    """else: # our target is not in it so we add it unchanged
                        paragraph.runs.append(run)"""

    doc.save("app_output\\pronouns_highlight_output.docx")
    os.system("start app_output\\pronouns_highlight_output.docx")

coding("app_input\\test_pronoun.docx")

